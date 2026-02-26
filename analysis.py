"""
analysis.py — Reproducible artifact generator for the CFM56-7B FMEA–MCDM case study.

What it generates (local run):
- outputs/dataset_and_results.xlsx
- outputs/figures/Figure1_Workflow.png
- outputs/figures/Figure2_RankingComparison.png
- outputs/figures/Figure3_CRITIC_Weights.png
- outputs/figures/Figure4_Sensitivity_wS.png
Optional:
- outputs/01_FMEA_manuscript_UPDATED.docx (if manuscript docx exists locally)
- outputs/SupplementaryMaterial_UPDATED.docx (if supplementary docx exists locally)

Run locally:
    pip install -r requirements.txt
    python analysis.py

Data sources (conceptual):
- Occurrence index O: derived from FAA SDR keyword-count proxies and normalized by the maximum category count.
- Severity S and Detection D: defined on 1–10 ordinal scales grounded in NTSB investigation evidence and inspection detectability logic.

NOTE:
This repository is designed for transparency and reproducibility. Exact SDR query settings and the two NTSB reports
are described in the manuscript and supplementary material.
"""

from __future__ import annotations

import os
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

# Optional (only needed if you want to auto-update .docx locally)
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    Document = None


# -----------------------------
# CONFIG — EDIT HERE IF NEEDED
# -----------------------------
FAILURE_MODES = ["Fan blade", "Oil system", "Fuel system"]

# Occurrence index O is normalized by max(f) = 253 (fuel-related keyword proxy count).
# Fan blade O uses "engine shutdown" proxy count 108 (see manuscript methodology).
RAW_MATRIX = pd.DataFrame(
    {
        "Failure Mode": FAILURE_MODES,
        "O": [0.427, 0.541, 1.000],
        "S": [10, 8, 7],
        "D": [6, 4, 6],
    }
)

OUTDIR = "outputs"
FIGDIR = os.path.join(OUTDIR, "figures")

MANUSCRIPT_IN = "01 FMEA makale -v2.docx"
SUPP_IN = "Supplementary Material.docx"


# -----------------------------
# CORE METHODS
# -----------------------------
def vector_normalize(df: pd.DataFrame, cols: list[str]) -> pd.DataFrame:
    out = df.copy()
    for c in cols:
        denom = np.sqrt((df[c].astype(float) ** 2).sum())
        out[c] = df[c].astype(float) / denom
    return out


def critic_weights_raw(df: pd.DataFrame, cols: list[str]):
    """
    CRITIC weights computed on the raw matrix.
    - O is already a normalized occurrence index
    - S and D are 1–10 ordinal scores
    This design is explicitly documented in the manuscript and Supplementary tables.
    """
    X = df[cols].astype(float).values
    std = X.std(axis=0, ddof=0)  # population std
    corr = np.corrcoef(X.T)

    C = []
    for j in range(len(cols)):
        Cj = std[j] * np.sum(1.0 - corr[j])
        C.append(Cj)

    C = np.array(C)
    w = C / C.sum()
    return std, corr, C, w


def topsis(df: pd.DataFrame, cols: list[str], w: np.ndarray) -> pd.DataFrame:
    """
    TOPSIS on vector-normalized matrix, with benefit-type criteria (higher = more critical risk).
    """
    X = df[cols].astype(float).values
    R = X / np.sqrt((X**2).sum(axis=0))
    V = R * w

    ideal = V.max(axis=0)
    nadir = V.min(axis=0)

    d_plus = np.sqrt(((V - ideal) ** 2).sum(axis=1))
    d_minus = np.sqrt(((V - nadir) ** 2).sum(axis=1))
    score = d_minus / (d_plus + d_minus)

    return (
        pd.DataFrame(
            {"Failure Mode": df["Failure Mode"], "D_plus": d_plus, "D_minus": d_minus, "Score": score}
        )
        .sort_values("Score", ascending=False)
        .reset_index(drop=True)
    )


def vikor(df: pd.DataFrame, cols: list[str], w: np.ndarray, v: float = 0.5) -> pd.DataFrame:
    """
    VIKOR on vector-normalized matrix (benefit-type).
    Lower Q = more critical (closer to best compromise).
    """
    X = df[cols].astype(float).values
    R = X / np.sqrt((X**2).sum(axis=0))

    f_star = R.max(axis=0)
    f_minus = R.min(axis=0)
    denom = np.where((f_star - f_minus) == 0, 1e-12, (f_star - f_minus))

    S = np.sum(w * (f_star - R) / denom, axis=1)
    Rr = np.max(w * (f_star - R) / denom, axis=1)

    S_star, S_minus = S.min(), S.max()
    R_star, R_minus = Rr.min(), Rr.max()

    Q = v * (S - S_star) / (S_minus - S_star + 1e-12) + (1 - v) * (Rr - R_star) / (R_minus - R_star + 1e-12)

    return (
        pd.DataFrame({"Failure Mode": df["Failure Mode"], "S": S, "R": Rr, "Q": Q})
        .sort_values("Q", ascending=True)
        .reset_index(drop=True)
    )


def sensitivity_grid(raw: pd.DataFrame, cols: list[str], w_base: np.ndarray) -> pd.DataFrame:
    """
    Simple sensitivity: multiply severity weight pre-renormalization and recompute TOPSIS order.
    """
    def renorm(w):
        w = np.array(w, dtype=float)
        return w / w.sum()

    grid = np.linspace(0.6, 1.4, 17)
    rows = []
    for m in grid:
        w2 = w_base.copy()
        w2[1] = w_base[1] * m  # severity index is 1
        w2 = renorm(w2)
        order = list(topsis(raw, cols, w2)["Failure Mode"])
        rows.append(
            {"mult_S": m, "wO": w2[0], "wS": w2[1], "wD": w2[2], "TOPSIS_top": order[0], "TOPSIS_order": " > ".join(order)}
        )
    return pd.DataFrame(rows)


# -----------------------------
# OUTPUT HELPERS
# -----------------------------
def ensure_dirs():
    os.makedirs(OUTDIR, exist_ok=True)
    os.makedirs(FIGDIR, exist_ok=True)


def save_excel(raw: pd.DataFrame, cols: list[str], std, corr, C, w, topsis_df, vikor_df, sens_df):
    path = os.path.join(OUTDIR, "dataset_and_results.xlsx")

    raw2 = raw.copy()
    raw2["RPN"] = raw2["O"] * raw2["S"] * raw2["D"]
    rpn_rank = raw2[["Failure Mode", "RPN"]].sort_values("RPN", ascending=False).reset_index(drop=True)

    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        raw2.to_excel(writer, sheet_name="RawMatrix", index=False)
        vector_normalize(raw, cols).to_excel(writer, sheet_name="NormMatrix_Vector", index=False)
        rpn_rank.to_excel(writer, sheet_name="RPN", index=False)
        pd.DataFrame({"Criterion": cols, "StdDev_raw": std, "Cj": C, "Weight": w}).to_excel(writer, sheet_name="CRITIC", index=False)
        pd.DataFrame(corr, index=cols, columns=cols).to_excel(writer, sheet_name="CRITIC_Corr")
        topsis_df.to_excel(writer, sheet_name="TOPSIS", index=False)
        vikor_df.to_excel(writer, sheet_name="VIKOR", index=False)
        sens_df.to_excel(writer, sheet_name="Sensitivity_S", index=False)


def make_figures(raw: pd.DataFrame, cols: list[str], w: np.ndarray, sens_df: pd.DataFrame):
    # Figure 1 — workflow
    plt.figure(figsize=(10, 4))
    ax = plt.gca()
    ax.axis("off")
    boxes = [
        ("Failure data\n(NTSB + FAA SDR)", 0.05, 0.55),
        ("Data-driven FMEA\n(O,S,D construction)", 0.32, 0.55),
        ("Objective weighting\n(CRITIC)", 0.60, 0.55),
        ("Robust ranking\n(TOPSIS + VIKOR)", 0.83, 0.55),
    ]
    for text, x, y in boxes:
        rect = plt.Rectangle((x, y), 0.22, 0.30, fill=False, linewidth=2)
        ax.add_patch(rect)
        ax.text(x + 0.11, y + 0.15, text, ha="center", va="center", fontsize=10)
    for i in range(len(boxes) - 1):
        x1 = boxes[i][1] + 0.22
        y1 = boxes[i][2] + 0.15
        x2 = boxes[i + 1][1]
        y2 = boxes[i + 1][2] + 0.15
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1), arrowprops=dict(arrowstyle="->", lw=2))
    plt.tight_layout()
    plt.savefig(os.path.join(FIGDIR, "Figure1_Workflow.png"), dpi=300, bbox_inches="tight")
    plt.close()

    # Figure 2 — ranking comparison (RPN vs TOPSIS vs VIKOR)
    topsis_df = topsis(raw, cols, w)
    vikor_df = vikor(raw, cols, w)

    raw2 = raw.copy()
    raw2["RPN"] = raw2["O"] * raw2["S"] * raw2["D"]
    rpn_order = list(raw2.sort_values("RPN", ascending=False)["Failure Mode"])
    topsis_order = list(topsis_df["Failure Mode"])
    vikor_order = list(vikor_df["Failure Mode"])

    methods = ["RPN", "TOPSIS", "VIKOR"]
    ranks = pd.DataFrame(index=list(raw["Failure Mode"]), columns=methods, dtype=int)
    for m, order in zip(methods, [rpn_order, topsis_order, vikor_order]):
        for i, fm in enumerate(order, start=1):
            ranks.loc[fm, m] = i

    plt.figure(figsize=(7, 4))
    ax = plt.gca()
    x = np.arange(len(ranks.index))
    width = 0.25
    for i, m in enumerate(methods):
        ax.bar(x + i * width, ranks[m].values, width, label=m)
    ax.set_xticks(x + width)
    ax.set_xticklabels(ranks.index)
    ax.set_ylabel("Rank (1 = most critical)")
    ax.invert_yaxis()
    ax.legend()
    plt.tight_layout()
    plt.savefig(os.path.join(FIGDIR, "Figure2_RankingComparison.png"), dpi=300, bbox_inches="tight")
    plt.close()

    # Figure 3 — CRITIC weights
    plt.figure(figsize=(6, 4))
    ax = plt.gca()
    ax.bar(cols, w)
    ax.set_ylabel("Weight")
    ax.set_title("CRITIC objective weights")
    plt.tight_layout()
    plt.savefig(os.path.join(FIGDIR, "Figure3_CRITIC_Weights.png"), dpi=300, bbox_inches="tight")
    plt.close()

    # Figure 4 — sensitivity curve (how normalized wS changes after reweighting)
    plt.figure(figsize=(7, 4))
    ax = plt.gca()
    ax.plot(sens_df["mult_S"], sens_df["wS"], marker="o")
    ax.set_xlabel("Severity weight multiplier (before renormalization)")
    ax.set_ylabel("Final normalized w_S")
    ax.set_title("Sensitivity of w_S under reweighting")
    plt.tight_layout()
    plt.savefig(os.path.join(FIGDIR, "Figure4_Sensitivity_wS.png"), dpi=300, bbox_inches="tight")
    plt.close()


def try_update_docx():
    """
    Optional: if .docx files exist locally, insert figures after captions.
    This does nothing on GitHub web (no runtime), but supports local regeneration.
    """
    if Document is None:
        return

    if not os.path.exists(MANUSCRIPT_IN) or not os.path.exists(SUPP_IN):
        return

    def insert_image_after_caption(doc, caption_startswith, image_path, width_inches=6.0):
        for p in doc.paragraphs:
            if p.text.strip().startswith(caption_startswith):
                new_p = doc.add_paragraph()
                p._p.addnext(new_p._p)
                run = new_p.add_run()
                run.add_picture(image_path, width=Inches(width_inches))
                new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                return True
        return False

    man = Document(MANUSCRIPT_IN)
    insert_image_after_caption(man, "Figure 1.", os.path.join(FIGDIR, "Figure1_Workflow.png"), 6.5)
    insert_image_after_caption(man, "Figure 2.", os.path.join(FIGDIR, "Figure2_RankingComparison.png"), 6.0)
    man.save(os.path.join(OUTDIR, "01_FMEA_manuscript_UPDATED.docx"))


def main():
    ensure_dirs()
    cols = ["O", "S", "D"]

    std, corr, C, w = critic_weights_raw(RAW_MATRIX, cols)
    topsis_df = topsis(RAW_MATRIX, cols, w)
    vikor_df = vikor(RAW_MATRIX, cols, w)
    sens_df = sensitivity_grid(RAW_MATRIX, cols, w)

    save_excel(RAW_MATRIX, cols, std, corr, C, w, topsis_df, vikor_df, sens_df)
    make_figures(RAW_MATRIX, cols, w, sens_df)

    try_update_docx()

    print("Done. See outputs/ for Excel and figures.")


if __name__ == "__main__":
    main()
